"""Extract text from uploaded documents (PDF, DOCX, TXT, CSV, MD, JSON).

Used to provide document context to the AI when users upload files.
"""

import base64
import io
import logging

log = logging.getLogger(__name__)

# Mime types that Gemini can handle natively as binary
GEMINI_NATIVE_TYPES = {
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/webp",
}

# Mime types we can extract text from
TEXT_MIME_TYPES = {
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
    "application/json",
    "application/xml",
    "text/xml",
}

# Extensions that are text-based
TEXT_EXTENSIONS = {".txt", ".csv", ".md", ".json", ".xml", ".html", ".py", ".js", ".ts", ".yaml", ".yml", ".toml"}


def is_text_file(mime_type: str, filename: str = "") -> bool:
    """Check if a file can be read as text."""
    if mime_type in TEXT_MIME_TYPES:
        return True
    ext = _get_ext(filename)
    return ext in TEXT_EXTENSIONS


def is_gemini_native(mime_type: str) -> bool:
    """Check if Gemini can handle this file type natively as binary."""
    return mime_type in GEMINI_NATIVE_TYPES or mime_type.startswith("image/")


def _get_ext(filename: str) -> str:
    """Get lowercase file extension."""
    if "." in filename:
        return "." + filename.rsplit(".", 1)[-1].lower()
    return ""


async def extract_text(base64_data: str, mime_type: str, filename: str = "") -> str:
    """Extract text content from a document.

    Returns extracted text, or empty string if the file should be
    sent as binary to Gemini (e.g., PDF, images).
    """
    try:
        raw = base64.b64decode(base64_data)
    except Exception:
        log.exception("Failed to decode base64 document data")
        return ""

    # Text-based files — decode directly
    if is_text_file(mime_type, filename):
        try:
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return raw.decode("latin-1", errors="replace")

    # DOCX — parse with python-docx
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or _get_ext(filename) in (".docx", ".doc"):
        return _extract_docx(raw)

    # PDF — try to extract text; if that fails, Gemini can handle binary
    if mime_type == "application/pdf" or _get_ext(filename) == ".pdf":
        text = _extract_pdf(raw)
        if text.strip():
            return text
        # Return empty — will be sent as binary to Gemini
        return ""

    # Unknown type — try as text
    try:
        return raw.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())
        return "\n\n".join(text_parts)
    except ImportError:
        log.warning("pypdf not installed — PDF will be sent as binary to Gemini")
        return ""
    except Exception:
        log.exception("PDF text extraction failed")
        return ""


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        return "\n\n".join(text_parts)
    except ImportError:
        log.warning("python-docx not installed — cannot parse DOCX")
        return ""
    except Exception:
        log.exception("DOCX text extraction failed")
        return ""
